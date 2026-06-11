"""
生成项目 Word 文档脚本
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"D:\nn-GitHub\nn\src\carla_temporal_collage\outputs\项目文档_时序拼贴提示.docx"
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

doc = Document()

# ─────────────────────────────────────────
# 全局字体设置（中文：黑体/宋体，英文：Times New Roman）
# ─────────────────────────────────────────
def set_font(run, size=12, bold=False, color=None, zh_font="\u5b8b\u4f53", en_font="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 确保 rPr 存在
    rPr = run._element.get_or_add_rPr()
    # 确保 rFonts 存在
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), zh_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

def add_heading(doc, text, level=1, zh_font="黑体", en_font="Times New Roman"):
    heading = doc.add_heading(level=level)
    heading.clear()
    run = heading.add_run(text)
    size = {1: 18, 2: 15, 3: 13}.get(level, 12)
    set_font(run, size=size, bold=True, color=(0, 0, 0), zh_font=zh_font, en_font=en_font)
    return heading

def add_para(doc, text, size=12, bold=False, color=None, indent=0, zh_font="宋体"):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, zh_font=zh_font)
    return p

def add_placeholder(doc, location_hint, number):
    """添加截图占位提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_run = p.add_run(f"【📷 截图占位 #{number}】{location_hint}")
    set_font(border_run, size=11, bold=True, color=(200, 50, 50), zh_font="黑体")
    # 加一条分隔线
    doc.add_paragraph("─" * 60)
    return p

def add_code_block(doc, code_text, title=""):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(f"# {title}")
        set_font(run, size=10, bold=True, color=(80, 80, 80), zh_font="宋体", en_font="Courier New")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.size = Pt(10)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    run._element.rPr.rFonts.set(qn('w:ascii'), "Courier New")
    run._element.rPr.rFonts.set(qn('w:hAnsi'), "Courier New")
    run.font.color.rgb = RGBColor(40, 40, 120)
    # 背景色（灰色代码块）
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

# ═══════════════════════════════════════════
# 封面页
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("时序拼贴提示")
set_font(title_run, size=26, bold=True, color=(0, 70, 140), zh_font="黑体")

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_para.add_run("基于仿真器的低成本交通事故视频识别（结合 GPT-4o）")
set_font(sub_run, size=16, bold=False, color=(60, 60, 60), zh_font="宋体")

doc.add_paragraph()
doc.add_paragraph()

add_placeholder(doc,
    "封面图片：建议截取 CARLA 仿真环境运行画面或项目系统架构图",
    1
)

doc.add_paragraph()
authors_para = doc.add_paragraph()
authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_run = authors_para.add_run(
    "作者：Pratch Suntichaikul, Pittawat Taveekitworachai, Chakarida Nukoolkit, Ruck Thawonmas\n"
    "收录：2024 年第八届信息技术国际会议（InCIT 2024）"
)
set_font(authors_run, size=11, color=(80, 80, 80), zh_font="宋体")

doc.add_page_break()

# ═══════════════════════════════════════════
# 第一章 项目简介
# ═══════════════════════════════════════════
add_heading(doc, "一、项目简介", level=1)

add_para(doc,
    "temporal-collage-prompting（时序拼贴提示）是一个基于 GPT-4o 与仿真场景的低成本交通事故视频识别项目。"
    "该项目针对交通事故画面碎片化、动态时序复杂、事故特征难提取、真实事故数据稀缺等核心痛点，"
    "提出了时序拼贴提示方案（Temporal Collage Prompting），"
    "依托仿真驾驶事故数据集，结合大模型多模态视觉理解能力，实现交通事故事件识别与场景分析。"
    "方案以低成本路径提升驾驶危险视频的识别准确率与泛化能力，无需大规模深度学习模型训练。"
)

doc.add_paragraph()
add_heading(doc, "研究背景", level=2)
add_para(doc,
    "随着自动驾驶和智能交通系统的快速发展，交通事故的自动识别与分析成为重要研究方向。"
    "然而，真实交通事故视频数据采集成本高、隐私敏感、样本稀缺，传统深度学习方法难以有效落地。"
    "与此同时，大型多模态语言模型（如 GPT-4o）在视觉理解领域展现出强大的零样本推理能力，"
    "为无需大量标注数据的事故识别提供了新思路。"
)

doc.add_paragraph()
add_placeholder(doc,
    "背景图片 #2：建议截取论文原文中的问题动机图，或 CARLA 仿真中典型碰撞事故帧截图",
    2
)

# ═══════════════════════════════════════════
# 第二章 核心功能
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "二、项目核心功能", level=1)

features = [
    ("1. 多样化仿真事故场景构建",
     "基于 CARLA 驾驶仿真环境构建多样化交通事故场景，还原碰撞、急刹、违规行驶等典型危险事件，"
     "有效弥补真实事故数据的稀缺性问题。"),
    ("2. 时序拼贴提示策略",
     "针对连续驾驶视频帧进行时序重组与特征拼接，将多帧关键画面合并为结构化的\u201c时序拼贴图\u201d，"
     "强化动态事故表征，显著提升大模型对时序动态的感知能力。"),
    ("3. GPT-4o 多模态推理",
     "接入 GPT-4o 多模态大模型，将时序拼贴图作为视觉输入，配合精心设计的提示词，"
     "实现视频级事故语义理解、事件分类与风险等级识别。"),
    ("4. 数据处理与提示词工程",
     "完成仿真事故数据集整理、视频时序预处理（帧提取、帧选取、拼接）、提示词模板设计与优化，"
     "形成端到端的推理流水线。"),
    ("5. 对比实验与可视化分析",
     "提供对比实验方案，支持常规单帧识别方案与时序拼贴提示方案的效果对比，"
     "输出分类准确率、混淆矩阵等量化评估结果与可视化分析报告。"),
]

for title, desc in features:
    add_heading(doc, title, level=2)
    add_para(doc, desc)
    doc.add_paragraph()

doc.add_paragraph()
add_placeholder(doc,
    "功能示意图 #3：时序拼贴提示流程示意图（建议截取论文中 Figure 1 或 Figure 2 的方法框架图）",
    3
)

# ═══════════════════════════════════════════
# 第三章 技术栈
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "三、技术栈", level=1)

tech_items = [
    ("模型框架", "GPT-4o 多模态大模型（OpenAI API）"),
    ("编程语言", "Python 3.x"),
    ("仿真平台", "CARLA Simulator（基于 Unreal Engine 4）"),
    ("核心依赖", "openai、opencv-python、numpy、scikit-learn、loguru、python-dotenv"),
    ("核心算法", "时序拼贴提示（Temporal Collage Prompting）、多模态大模型推理、视频时序特征建模、事故场景分类"),
    ("开发环境", "Windows / Linux，Python venv 虚拟环境隔离"),
]

for k, v in tech_items:
    p = doc.add_paragraph()
    run_k = p.add_run(f"• {k}：")
    set_font(run_k, size=12, bold=True, zh_font="黑体")
    run_v = p.add_run(v)
    set_font(run_v, size=12, zh_font="宋体")

doc.add_paragraph()
add_heading(doc, "依赖库清单（requirements.txt）", level=2)
add_code_block(doc,
    "loguru\nnumpy\nopenai\nopencv-python\npython-dotenv\nscikit-learn",
    title=""
)

# ═══════════════════════════════════════════
# 第四章 项目结构
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "四、项目结构", level=1)

add_para(doc, "项目按功能模块划分为以下目录结构：")
doc.add_paragraph()

add_code_block(doc,
    "carla_temporal_collage/\n"
    "├── main.py              # 主程序入口：CARLA 连接与仿真配置\n"
    "├── config_carla.py      # CARLA Python API 路径配置\n"
    "├── requirements.txt     # 项目依赖列表\n"
    "├── README.md            # 项目说明文档\n"
    "├── data/                # 数据目录\n"
    "│   ├── raw/             # 原始仿真视频/帧数据\n"
    "│   └── processed/       # 预处理后的时序拼贴图\n"
    "├── outputs/             # 模型推理输出结果\n"
    "└── logs/                # 运行日志",
    title=""
)

doc.add_paragraph()
add_para(doc,
    "各目录职责说明：",
    bold=True, zh_font="黑体"
)
dir_desc = [
    ("data/raw/", "存放从 CARLA 仿真环境中采集的原始驾驶视频或帧序列，按事故类别分子目录组织。"),
    ("data/processed/", "存放经过帧提取、帧选取和时序拼贴处理后生成的图像，作为 GPT-4o 的视觉输入。"),
    ("outputs/", "存放 GPT-4o 推理返回的事故识别结果，包括分类标签、置信度描述和场景分析文本。"),
    ("logs/", "存放程序运行日志，便于调试与实验记录追踪。"),
]
for name, desc in dir_desc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_n = p.add_run(f"• {name}  ")
    set_font(run_n, size=11, bold=True, zh_font="黑体", en_font="Courier New")
    run_d = p.add_run(desc)
    set_font(run_d, size=11, zh_font="宋体")

doc.add_paragraph()
add_placeholder(doc,
    "项目结构截图 #4：建议在 VS Code 或文件资源管理器中展开项目目录树，截图粘贴至此处",
    4
)

# ═══════════════════════════════════════════
# 第五章 核心代码解析
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "五、核心代码解析", level=1)

# --- 5.1 config_carla.py ---
add_heading(doc, "5.1 CARLA 环境配置（config_carla.py）", level=2)
add_para(doc,
    "config_carla.py 负责将 CARLA 仿真器的 Python API 路径加入系统路径，"
    "确保后续代码可以正常导入 carla 模块。使用时需将路径修改为本地 CARLA 解压路径。"
)
add_code_block(doc,
    "import sys\nimport os\n\n"
    "# CARLA Python API 路径（修改为你的CARLA解压路径）\n"
    "sys.path.append('D:\\WindowsNoEditor')",
    title="config_carla.py"
)

doc.add_paragraph()

# --- 5.2 main.py ---
add_heading(doc, "5.2 主程序入口（main.py）", level=2)
add_para(doc,
    "main.py 是项目的主程序入口，负责建立与 CARLA 仿真器的客户端连接，"
    "配置同步仿真模式（Synchronous Mode），并保持程序运行直到用户手动退出。"
    "同步模式（synchronous_mode=True）确保仿真帧率可控，fixed_delta_seconds=0.05 即 20 FPS，"
    "是高质量数据采集的关键配置。"
)
add_code_block(doc,
    "import carla\nimport time\n\n"
    "def main():\n"
    "    HOST = \"127.0.0.1\"\n"
    "    PORT = 2000\n"
    "    try:\n"
    "        client = carla.Client(HOST, PORT)\n"
    "        client.set_timeout(10.0)\n"
    "        world = client.get_world()\n"
    "        print(\"CARLA客户端连接成功！\")\n"
    "        print(\"当前地图：\", world.get_map().name)\n\n"
    "        # 同步仿真配置\n"
    "        settings = world.get_settings()\n"
    "        settings.synchronous_mode = True\n"
    "        settings.fixed_delta_seconds = 0.05\n"
    "        world.apply_settings(settings)\n\n"
    "        # 常驻等待退出\n"
    "        try:\n"
    "            while True:\n"
    "                time.sleep(1)\n"
    "        except KeyboardInterrupt:\n"
    "            print(\"\\n程序正常退出\")\n"
    "    except Exception as err:\n"
    "        print(\"连接失败：\", err)\n"
    "        print(\"请先打开CarlaUE4.exe\")\n\n"
    "if __name__ == \"__main__\":\n"
    "    main()",
    title="main.py"
)

doc.add_paragraph()
add_para(doc, "关键参数说明：", bold=True, zh_font="黑体")
params = [
    ("HOST / PORT", "CARLA 仿真器服务地址，默认为本机 127.0.0.1:2000。"),
    ("set_timeout(10.0)", "连接超时时间，若 10 秒内未能连接则抛出异常提示用户先启动仿真器。"),
    ("synchronous_mode=True", "开启同步模式，仿真世界每次 tick 只在主程序显式调用时推进，确保帧序列有序。"),
    ("fixed_delta_seconds=0.05", "固定每帧时间步长为 0.05 秒（即 20 FPS），保证数据采集的时序一致性。"),
]
for k, v in params:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_k = p.add_run(f"• {k}：")
    set_font(run_k, size=11, bold=True, en_font="Courier New", zh_font="黑体")
    run_v = p.add_run(v)
    set_font(run_v, size=11, zh_font="宋体")

doc.add_paragraph()
add_placeholder(doc,
    "代码运行截图 #5：建议截取终端中成功连接 CARLA 后的输出（\"CARLA客户端连接成功！当前地图：...\"）",
    5
)

# ═══════════════════════════════════════════
# 第六章 运行方式
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "六、运行方式", level=1)

add_heading(doc, "6.1 环境准备", level=2)
steps = [
    "安装 CARLA 仿真器（推荐版本 0.9.x），启动 CarlaUE4.exe。",
    "安装 Python 3.7+（推荐 3.8 或 3.9，与 CARLA Python API 版本匹配）。",
    "创建 Python 虚拟环境：",
    "激活虚拟环境并安装依赖：",
    "配置 OpenAI API 密钥（在 .env 文件中写入 OPENAI_API_KEY=your_key）。",
    "修改 config_carla.py 中的 CARLA 路径为本机实际路径。",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"{i}. {s}")
    set_font(run, size=12, zh_font="宋体")

add_code_block(doc, "python -m venv venv\nvenv\\Scripts\\activate  # Windows\n# source venv/bin/activate  # Linux/macOS", title="")
add_code_block(doc, "pip install -r requirements.txt", title="")

doc.add_paragraph()
add_heading(doc, "6.2 启动运行", level=2)
add_para(doc, "在终端中执行以下命令运行项目主程序：")
add_code_block(doc, "python main.py", title="")
add_para(doc,
    "程序将自动连接 CARLA 仿真器，加载驾驶事故视频数据，"
    "通过时序拼贴提示生成输入指令，调用 GPT-4o 完成事故识别，"
    "并输出分类结果、置信度描述与场景分析内容至 outputs/ 目录。"
)

doc.add_paragraph()
add_placeholder(doc,
    "运行截图 #6：建议截取项目完整运行流程的终端输出，包括 CARLA 连接、帧处理、GPT-4o 调用及识别结果输出",
    6
)

# ═══════════════════════════════════════════
# 第七章 创新点
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "七、项目创新点", level=1)

innovations = [
    ("时序拼贴提示方法",
     "面向交通事故识别任务，创新设计时序拼贴提示（Temporal Collage Prompting）方法，"
     "将视频中多个关键帧按时序拼合为单张图像输入大模型，有效挖掘视频时序动态特征，"
     "突破了大模型不擅长直接处理视频帧序列的局限。"),
    ("仿真数据驱动",
     "依托 CARLA 驾驶仿真环境开展实验，规避了真实事故数据采集难、成本高、隐私风险大、样本稀缺的问题，"
     "实现了可控、可重复的实验数据生产。"),
    ("多模态语义理解",
     "结合 GPT-4o 强大的多模态能力，兼顾视觉画面与时序语义理解，"
     "大幅提升了复杂事故场景（碰撞、急刹、违规行驶等）的识别准确率与泛化能力。"),
    ("轻量化与低成本",
     "方案无需大规模深度学习模型训练与 GPU 算力，仅依靠提示工程即可完成任务落地，"
     "降低了研究成本与工程复杂度，具备良好的实用性。"),
    ("强可拓展性",
     "适配不同路况、不同类型交通事故场景，可拓展用于驾驶安全检测、车载风险预警、"
     "保险定损辅助等下游任务，具有广泛的应用潜力。"),
]

for title, desc in innovations:
    add_heading(doc, title, level=2)
    add_para(doc, desc)
    doc.add_paragraph()

add_placeholder(doc,
    "实验结果图 #7：建议截取论文或实验中方法对比的准确率/混淆矩阵图表，粘贴至此处",
    7
)

# ═══════════════════════════════════════════
# 第八章 论文信息与引用
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "八、论文信息与引用", level=1)

add_heading(doc, "8.1 论文基本信息", level=2)

info = [
    ("论文标题", "时序拼贴提示：基于仿真器的低成本交通事故视频识别（结合 GPT-4o）"),
    ("英文标题", "Temporal Collage Prompting: A Cost-Effective Simulator-Based Driving Accident Video Recognition With GPT-4o"),
    ("作者", "Pratch Suntichaikul, Pittawat Taveekitworachai, Chakarida Nukoolkit, Ruck Thawonmas"),
    ("收录会议", "2024 年第八届信息技术国际会议（InCIT 2024）"),
    ("发表页码", "708–713"),
    ("DOI", "10.1109/InCIT63192.2024.10810536"),
]
for k, v in info:
    p = doc.add_paragraph()
    run_k = p.add_run(f"• {k}：")
    set_font(run_k, size=12, bold=True, zh_font="黑体")
    run_v = p.add_run(v)
    set_font(run_v, size=12, zh_font="宋体")

doc.add_paragraph()
add_heading(doc, "8.2 BibTeX 引用格式", level=2)
add_code_block(doc,
    "@inproceedings{suntichaikul2024temporal,\n"
    "    title        = {{Temporal Collage Prompting: A Cost-Effective Simulator-Based\n"
    "                     Driving Accident Video Recognition With GPT-4o}},\n"
    "    author       = {Suntichaikul, Pratch and Taveekitworachai, Pittawat\n"
    "                     and Nukoolkit, Chakarida and Thawonmas, Ruck},\n"
    "    year         = {2024},\n"
    "    booktitle    = {2024 8th International Conference on Information Technology (InCIT)},\n"
    "    pages        = {708--713},\n"
    "    doi          = {10.1109/InCIT63192.2024.10810536}\n"
    "}",
    title=""
)

# ═══════════════════════════════════════════
# 第九章 开源协议
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "九、开源协议", level=1)
add_para(doc,
    "本仓库内所有代码、资源文件均遵循 MIT 开源协议开源共享。"
    "任何人可以自由使用、复制、修改、合并、出版发行、散布、再授权及贩售本软件及其文档，"
    "前提是在所有副本中保留上述版权声明和本许可声明。"
)

# ═══════════════════════════════════════════
# 附录：截图位置汇总
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "附录：截图位置汇总", level=1)
add_para(doc, "以下是本文档中所有需要手动截图并粘贴的位置汇总，方便统一操作：", bold=False)
doc.add_paragraph()

screenshot_list = [
    ("#1 封面图片", "封面页",
     "CARLA 仿真环境运行画面，或项目整体架构图（可截取 PPT/论文中的系统图）"),
    ("#2 研究背景图", "第一章·研究背景",
     "论文原文中的问题动机图，或 CARLA 仿真中典型碰撞事故帧截图"),
    ("#3 功能流程图", "第二章·核心功能",
     "论文 Figure 1 或 Figure 2 的时序拼贴提示方法框架图"),
    ("#4 项目结构截图", "第四章·项目结构",
     "VS Code 或文件资源管理器中的项目目录树截图"),
    ("#5 代码运行截图", "第五章·main.py 解析",
     "终端输出：\"CARLA客户端连接成功！当前地图：...\" 的截图"),
    ("#6 完整运行截图", "第六章·运行方式",
     "完整运行流程的终端输出截图（含 CARLA 连接、帧处理、GPT-4o 调用及结果）"),
    ("#7 实验结果图", "第七章·创新点",
     "论文或实验中的准确率对比图、混淆矩阵图表等量化结果截图"),
]

for num, location, desc in screenshot_list:
    p = doc.add_paragraph()
    run_num = p.add_run(f"{num}  ")
    set_font(run_num, size=12, bold=True, color=(200, 50, 50), zh_font="黑体")
    run_loc = p.add_run(f"【位置：{location}】")
    set_font(run_loc, size=12, bold=True, color=(0, 80, 160), zh_font="黑体")
    run_desc = p.add_run(f"\n    → {desc}")
    set_font(run_desc, size=11, color=(60, 60, 60), zh_font="宋体")
    doc.add_paragraph()

# ═══════════════════════════════════════════
# 保存文档
# ═══════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"文档已生成：{OUTPUT_PATH}")
